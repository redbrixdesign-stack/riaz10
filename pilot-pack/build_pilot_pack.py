from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_LINE_SPACING
from pathlib import Path

OUT = Path(__file__).resolve().parent / "Beelo-Pilot-Pack.docx"

GOLD = "FDB913"
GOLD_DARK = "9A6A00"
BLACK = "0A0A0A"
INK = "202020"
GRAY = "5F6368"
LIGHT = "F6F4ED"
PALE_GOLD = "FFF5D6"
PALE_RED = "FCE8E6"
RED = "A33A2B"
GREEN = "315B3A"
WHITE = "FFFFFF"
BORDER = "D8D4C8"
FONT = "Calibri"
WIDTH_DXA = 9360
INDENT_DXA = 120


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths, indent=INDENT_DXA):
    assert sum(widths) == WIDTH_DXA
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(width))
        grid.append(gc)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table, color=BORDER, size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:color"), color)


def keep_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_font(run, size=None, color=INK, bold=None, italic=None, name=FONT):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def style_para(p, before=0, after=6, line=1.25, keep=False):
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    p.paragraph_format.keep_together = keep


def add_para(doc, text="", size=10.5, color=INK, bold=False, italic=False,
             before=0, after=6, line=1.25, align=None, keep=False):
    p = doc.add_paragraph()
    style_para(p, before, after, line, keep)
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_font(r, size, color, bold, italic)
    return p


def add_label_para(doc, label, text, after=6, color=INK):
    p = doc.add_paragraph()
    style_para(p, 0, after, 1.25)
    r = p.add_run(label + "  ")
    set_font(r, 10.5, GOLD_DARK, True)
    r = p.add_run(text)
    set_font(r, 10.5, color)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    if getattr(doc, "_beelo_page_break_pending", False):
        p.paragraph_format.page_break_before = True
        doc._beelo_page_break_pending = False
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    return p


def add_kicker(doc, text):
    p = doc.add_paragraph()
    if getattr(doc, "_beelo_page_break_pending", False):
        p.paragraph_format.page_break_before = True
        doc._beelo_page_break_pending = False
    style_para(p, 0, 5, 1.0, True)
    r = p.add_run(text.upper())
    set_font(r, 9, GOLD_DARK, True)
    r.font.letter_spacing = Pt(1.2) if hasattr(r.font, "letter_spacing") else None
    return p


def add_callout(doc, label, text, fill=PALE_GOLD, accent=GOLD_DARK):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [WIDTH_DXA], indent=180)
    set_table_borders(table, fill, "0")
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, 160, 180, 160, 180)
    p = cell.paragraphs[0]
    style_para(p, 0, 0, 1.2)
    r = p.add_run(label + "  ")
    set_font(r, 10, accent, True)
    r = p.add_run(text)
    set_font(r, 10, INK)
    add_para(doc, "", size=2, after=4, line=1.0)
    return table


def add_two_col_facts(doc, rows):
    table = doc.add_table(rows=len(rows), cols=2)
    set_table_geometry(table, [2700, 6660])
    set_table_borders(table)
    for i, (label, value) in enumerate(rows):
        keep_row(table.rows[i])
        set_cell_shading(table.cell(i, 0), LIGHT)
        p = table.cell(i, 0).paragraphs[0]
        style_para(p, 0, 0, 1.15)
        set_font(p.add_run(label), 9.5, BLACK, True)
        p = table.cell(i, 1).paragraphs[0]
        style_para(p, 0, 0, 1.15)
        set_font(p.add_run(value), 9.5, INK)
    add_para(doc, "", size=2, after=4, line=1.0)
    return table


def add_matrix(doc, headers, rows, widths, font_size=8.7):
    table = doc.add_table(rows=1, cols=len(headers))
    for _ in rows:
        table.add_row()
    set_table_geometry(table, widths)
    set_table_borders(table)
    set_repeat_header(table.rows[0])
    for j, h in enumerate(headers):
        set_cell_shading(table.cell(0, j), BLACK)
        p = table.cell(0, j).paragraphs[0]
        style_para(p, 0, 0, 1.1)
        set_font(p.add_run(h), font_size, WHITE, True)
    for i, row in enumerate(rows, 1):
        keep_row(table.rows[i])
        for j, value in enumerate(row):
            p = table.cell(i, j).paragraphs[0]
            style_para(p, 0, 0, 1.12)
            set_font(p.add_run(str(value)), font_size, INK)
            if i % 2 == 0:
                set_cell_shading(table.cell(i, j), "FBFAF6")
    add_para(doc, "", size=2, after=4, line=1.0)
    return table


def add_checklist(doc, items, heading=None):
    if heading:
        add_heading(doc, heading, 3)
    table = doc.add_table(rows=len(items), cols=2)
    set_table_geometry(table, [500, 8860])
    set_table_borders(table, "E6E2D7")
    for i, item in enumerate(items):
        keep_row(table.rows[i])
        p = table.cell(i, 0).paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_para(p, 0, 0, 1.0)
        set_font(p.add_run("□"), 12, GOLD_DARK, True)
        p = table.cell(i, 1).paragraphs[0]
        style_para(p, 0, 0, 1.15)
        set_font(p.add_run(item), 9.5, INK)
    add_para(doc, "", size=2, after=4, line=1.0)
    return table


def add_field_line(doc, label, height_lines=1):
    p = doc.add_paragraph()
    style_para(p, 0, 8 if height_lines == 1 else 18, 1.2)
    set_font(p.add_run(label + "  "), 9.5, GRAY, True)
    set_font(p.add_run("____________________________________________________________"), 9.5, BORDER)
    if height_lines > 1:
        for _ in range(height_lines - 1):
            p.add_run("\n")
            set_font(p.add_run("             ____________________________________________________________"), 9.5, BORDER)


def page_break(doc):
    doc._beelo_page_break_pending = True


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)
    set_font(run, 8.5, GRAY)


def setup_styles(doc):
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    # Named preset override: slightly tighter vertical margins preserve usable
    # writing space in the printed field forms; side margins remain at 1 inch.
    sec.top_margin = Inches(0.8)
    sec.bottom_margin = Inches(0.8)
    sec.left_margin = Inches(1.0)
    sec.right_margin = Inches(1.0)
    sec.header_distance = Inches(0.35)
    sec.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    tokens = {
        1: (16, BLACK, 18, 10),
        2: (13, BLACK, 14, 7),
        3: (11.5, GOLD_DARK, 10, 5),
    }
    for level, (size, color, before, after) in tokens.items():
        s = doc.styles[f"Heading {level}"]
        s.font.name = FONT
        s._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        s._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = RGBColor.from_string(color)
        s.paragraph_format.space_before = Pt(before)
        s.paragraph_format.space_after = Pt(after)
        s.paragraph_format.line_spacing = 1.1
        s.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        s = doc.styles[style_name]
        s.font.name = FONT
        s.font.size = Pt(10.5)
        s.paragraph_format.left_indent = Inches(0.375)
        s.paragraph_format.first_line_indent = Inches(-0.188)
        s.paragraph_format.space_after = Pt(4)
        s.paragraph_format.line_spacing = 1.25


def setup_header_footer(doc):
    for sec in doc.sections:
        sec.different_first_page_header_footer = True
        hp = sec.header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        style_para(hp, 0, 0, 1.0)
        set_font(hp.add_run("BEELO  /  PILOT PACK"), 8.5, GRAY, True)
        fp = sec.footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        style_para(fp, 0, 0, 1.0)
        set_font(fp.add_run("Confidential working document  •  v1.0  •  "), 8.5, GRAY)
        add_page_field(fp)


def cover(doc):
    add_para(doc, "BEELO", size=15, color=BLACK, bold=True, after=92)
    add_kicker(doc, "Controlled UK pilot")
    add_para(doc, "Pilot Pack", size=34, color=BLACK, bold=True, after=8, line=1.0)
    add_para(doc, "A practical field guide for onboarding, running and evaluating a small pilot with solo field professionals.", size=15, color=GRAY, after=26, line=1.25)
    add_callout(doc, "POSITIONING", "Beelo is a device-local operational companion for the person doing the work: it helps hold together visits, customers, follow-ups, jobs, money records and user-approved communication drafts.")
    add_para(doc, "", size=3, after=42)
    add_two_col_facts(doc, [
        ("Prepared by", "Muhammad Asif Riaz, Founder of Beelo"),
        ("Contact", "hello@beelestial.co.uk"),
        ("Product", "beelo.beelestial.co.uk"),
        ("Pilot shape", "Recommended four weeks • 5–10 UK participants • controlled release"),
        ("Version", "v1.0 • 25 August 2026"),
    ])
    add_para(doc, "Internal and participant-facing working document. Pages marked REVIEW BEFORE ISSUE must be completed and reviewed before external use.", size=9, color=RED, bold=True, before=16, after=0)


def contents(doc):
    page_break(doc)
    add_kicker(doc, "How to use this pack")
    add_heading(doc, "One pack, three jobs", 1)
    add_para(doc, "Use the Founder pages to control the pilot, the Participant pages during onboarding, and the Research pages to collect comparable evidence. Duplicate the blank forms for each participant; keep the master copy unchanged.")
    add_matrix(doc, ["Part", "Use", "Primary reader"], [
        ("Pilot control", "Scope, readiness gates, claims and operating model", "Founder / partner"),
        ("Onboarding", "Setup, safe-use rules and participant acknowledgement", "Founder + participant"),
        ("Field operation", "Daily routine, demo flow and support", "Participant"),
        ("Evaluation", "Baseline, weekly pulse, interview and scorecard", "Founder / research partner"),
        ("Closeout", "Data choice, findings and go / iterate / stop decision", "Founder / partner"),
    ], [1600, 5060, 2700], 9)
    add_heading(doc, "Pilot principles", 2)
    for label, text in [
        ("Human controlled", "Beelo may organise, calculate or draft; the advisor reviews and decides. Nothing should be represented as silently sent or automatically decided."),
        ("Local by default", "Core working records are held in the participant's browser/device. There is no cloud sync, account recovery or automatic server backup in the current prototype."),
        ("Evidence before claims", "Admin-time savings, fewer missed follow-ups and reduced mental load are hypotheses to measure, not outcomes already achieved."),
        ("Small and reversible", "Begin with a controlled cohort, low-risk data and clear stop criteria. Expand only after the evidence and safeguards justify it."),
    ]:
        add_label_para(doc, label, text)
    add_callout(doc, "IMPORTANT", "This pack is operational guidance, not legal advice. The privacy and consent wording must be reviewed for the actual pilot arrangement, processors, lawful basis, retention periods and operator details.", PALE_RED, RED)


def pilot_at_glance(doc):
    page_break(doc)
    add_kicker(doc, "Founder / partner")
    add_heading(doc, "Pilot at a glance", 1)
    add_two_col_facts(doc, [
        ("Purpose", "Test whether Beelo is understandable, usable and helpful during real solo field work."),
        ("Recommended cohort", "5–10 UK-based solo or mainly solo field professionals; start with 3–5 if support capacity is limited."),
        ("Recommended duration", "Four active working weeks after onboarding."),
        ("Primary setting", "Home-visit window-coverings advisors; closely comparable field roles may be included deliberately."),
        ("Evidence level sought", "Independent target users complete core journeys and report outcomes (E4 pilot evidence)."),
        ("Pilot owner", "Muhammad Asif Riaz / Beelo."),
    ])
    add_heading(doc, "What is in scope now", 2)
    add_matrix(doc, ["Capability", "Pilot use", "Boundary"], [
        ("Visits and customers", "Plan visits, retain context and record outcomes", "No company diary replacement or automatic calendar sync"),
        ("Follow-ups", "Surface due work and create durable tasks", "Advisor remains responsible for action and timing"),
        ("Jobs and orders", "Link operational work to an order and track field status", "Does not process supplier orders or payments"),
        ("Money records", "Capture expenses, mileage and planning estimates", "Not bookkeeping, tax filing or accounting advice"),
        ("Quick Add / Scan", "Route a selected document or image into an intended workflow", "OCR quality varies; user verifies extracted data"),
        ("AI assistance", "Optional drafting, parsing and companion phrasing", "Requires connectivity; can be wrong; user reviews every output"),
        ("Navigation / timing", "Open the chosen map provider and record time on site", "External map estimates and OS behaviour are outside Beelo's control"),
    ], [1900, 3630, 3830], 8.4)
    page_break(doc)
    add_kicker(doc, "Founder / partner")
    add_heading(doc, "Pilot boundaries", 1)
    add_para(doc, "Use these boundaries consistently in demonstrations, recruitment, support and partner meetings. They protect the pilot from becoming a promise about a future product.")
    add_heading(doc, "Explicitly out of scope", 2)
    add_matrix(doc, ["Not in scope", "Boundary"], [
        ("Voice", "Not implemented; do not show or promise it."),
        ("Team platform", "Single-user and device-local; no manager dashboard or shared accounts."),
        ("Guaranteed outcomes", "Do not claim proven savings, scalability, compliance or certification."),
        ("Replacement system", "Does not replace a franchisor, supplier, CRM, accounting package or tax adviser."),
    ], [2050, 7310], 8.5)
    add_heading(doc, "Evidence language", 2)
    add_label_para(doc, "Implemented", "Use only where the current source and tests or direct observation support the capability.")
    add_label_para(doc, "Founder-tested", "Use for lived prototype experience; do not present it as independent user validation.")
    add_label_para(doc, "Pilot hypothesis", "Use for expected benefits such as reduced mental load, better follow-up reliability or time saved.")
    add_label_para(doc, "Unknown", "Say clearly when demand, willingness to pay, field reliability or scale has not yet been measured.")
    add_callout(doc, "MEETING CHECK", "If a statement would sound like an achieved customer outcome, measured result, certification or guarantee, pause and restate it as a capability boundary or pilot question.")


def hypotheses_and_success(doc):
    page_break(doc)
    add_kicker(doc, "Evaluation design")
    add_heading(doc, "Hypotheses and success measures", 1)
    add_para(doc, "The pilot should answer a small number of decision questions. Targets below are recommended starting thresholds, not performance claims or contractual service levels.")
    add_matrix(doc, ["Hypothesis", "Evidence", "Suggested success signal"], [
        ("Beelo reduces mental load", "Baseline/end rating and examples of remembered context", "At least 70% report a meaningful improvement"),
        ("Core journeys are learnable", "Observed onboarding and unassisted task completion", "At least 80% complete core setup and first visit flow"),
        ("It becomes part of the week", "Active days and meaningful records created", "Median 3+ active days per week after week one"),
        ("Human-controlled AI is trusted", "Draft review behaviour, corrections and interviews", "No silent-send expectation; users describe when they trust or reject it"),
        ("Local-first operation is acceptable", "Backup behaviour, device concerns and retention interviews", "Participants understand the trade-off and complete a safe backup"),
        ("The niche is worth pursuing", "Continued-use intent, referrals and willingness-to-pay discussion", "Clear repeat-use segment and specific unmet need emerge"),
    ], [2730, 3260, 3370], 8.5)
    add_heading(doc, "Decision rule", 2)
    add_matrix(doc, ["Decision", "Use when"], [
        ("GO", "Core workflows are repeatedly used; users report a specific benefit; no high-severity unresolved safety or data-loss issue; a credible next cohort exists."),
        ("ITERATE", "Need is real but onboarding, reliability, language or one core journey prevents repeat use. Fix a bounded problem and rerun."),
        ("STOP / PAUSE", "Users do not return, the problem is not important enough, safe operation cannot be maintained, or the support burden exceeds current capacity."),
    ], [1700, 7660], 9)
    add_callout(doc, "NON-NEGOTIABLE", "Zero unresolved high-severity privacy, unrecoverable-data-loss or unintended-communication incidents at the point of expansion.", PALE_RED, RED)


def readiness(doc):
    page_break(doc)
    add_kicker(doc, "Founder control")
    add_heading(doc, "Pre-pilot readiness gate", 1)
    add_para(doc, "Do not onboard external participants until every release-blocking item is complete or explicitly accepted with a documented workaround.")
    add_checklist(doc, [
        "Freeze a named pilot build and record commit, deployment, build date and service-worker cache version.",
        "Run the full automated suite, build check and core iPhone/Android browser journey against the pilot release.",
        "Remove test customers, mock visits, screenshots, expenses, API secrets and founder-only data from the participant setup path.",
        "Verify HTTPS, install, unlock, cold launch, supported offline shell, notification permission states and cache upgrade on physical devices.",
        "Verify the production AI proxy, allowed origin, rate limiting and shared-secret entry; define who receives the secret and how it is rotated.",
        "Complete operator name, address, contact email and company details shown in the product and privacy materials.",
        "Approve the pilot privacy notice, lawful basis, processor list, retention/deletion schedule and international-transfer explanation.",
        "Require encrypted backups or prohibit unsafe sharing/storage of readable backup files; test restore before onboarding.",
        "Prepare a support channel, issue log, incident response, stop criteria and a daily review owner.",
        "Test the pilot application, acknowledgement email and participant register end to end.",
        "Record supported device/browser versions and known limitations in the participant invitation.",
    ])
    add_heading(doc, "Release record", 2)
    add_field_line(doc, "Pilot release / commit")
    add_field_line(doc, "Deployment URL and date")
    add_field_line(doc, "Supported devices / browsers", 2)
    add_field_line(doc, "Accepted limitations / workarounds", 3)
    add_field_line(doc, "Approved by / date")


def roles(doc):
    page_break(doc)
    add_kicker(doc, "Operating model")
    add_heading(doc, "Roles, cadence and support", 1)
    add_matrix(doc, ["Role", "Responsibilities"], [
        ("Pilot lead", "Select participants, deliver onboarding, protect the release, review issues daily, maintain evidence and make stop/go decisions."),
        ("Participant", "Use Beelo only for agreed work, verify outputs, protect device/backups, report issues promptly and take part in agreed research."),
        ("Research / innovation partner", "Advise on method, reduce founder bias, support responsible-AI/privacy review and independently interpret evidence where agreed."),
        ("Technical support", "Reproduce defects, preserve user data, provide safe workarounds, release only controlled fixes and document changes."),
    ], [2200, 7160], 9)
    add_heading(doc, "Recommended cadence", 2)
    add_two_col_facts(doc, [
        ("Before start", "Eligibility call, baseline questionnaire and 35-minute onboarding."),
        ("Week one", "Two short check-ins; watch setup and first real journeys closely."),
        ("Weeks two–three", "Weekly pulse plus targeted interviews only where evidence is unclear."),
        ("Week four", "Final pulse, 30-minute exit interview, data choice and closeout."),
        ("After pilot", "Findings memo and go / iterate / pause review within five working days."),
    ])
    add_heading(doc, "Support details — complete before issue", 2)
    add_field_line(doc, "Support email")
    add_field_line(doc, "Support phone / WhatsApp")
    add_field_line(doc, "Normal support hours")
    add_field_line(doc, "Urgent data/privacy contact")
    add_callout(doc, "SUPPORT PROMISE", "Acknowledge problems clearly; do not promise immediate fixes. Protect participant and customer data before attempting recovery.")


def recruitment(doc):
    page_break(doc)
    add_kicker(doc, "Participant selection")
    add_heading(doc, "Recruitment and eligibility", 1)
    add_heading(doc, "Good-fit participant", 2)
    add_checklist(doc, [
        "UK-based and self-employed, or personally responsible for a meaningful share of field administration.",
        "Works alone or mainly alone across customer visits, travel and physical work.",
        "Uses several disconnected tools such as a company diary, WhatsApp, Maps, photos, notes, receipts or spreadsheets.",
        "Has a compatible personal device and is willing to install a PWA and protect it with a passphrase.",
        "Can use realistic but proportionate data and provide candid weekly feedback.",
        "Understands that this is an early controlled pilot, not a supported commercial service.",
    ])
    add_heading(doc, "Do not enrol yet", 2)
    add_checklist(doc, [
        "Anyone who requires team sharing, cross-device sync, guaranteed recovery or formal system-of-record status.",
        "Anyone planning to store special-category data, identity documents, payment-card details or confidential material they are not authorised to use.",
        "Anyone whose employer, franchise or supplier rules prohibit use of a personal operational tool.",
        "Anyone who cannot safely maintain a passphrase and encrypted backup.",
        "Anyone outside the UK for this first geographically controlled pilot.",
    ])
    add_heading(doc, "Screening record", 2)
    add_field_line(doc, "Participant name / ID")
    add_field_line(doc, "Trade / role and area")
    add_field_line(doc, "Current tool chain", 2)
    add_field_line(doc, "Biggest admin problem", 2)
    add_field_line(doc, "Eligibility outcome and reason", 2)


def onboarding(doc):
    page_break(doc)
    add_kicker(doc, "Participant session")
    add_heading(doc, "35-minute onboarding guide", 1)
    add_matrix(doc, ["Time", "Facilitator action", "Proof of understanding"], [
        ("0–5 min", "Explain the problem, pilot purpose, boundaries and voluntary nature.", "Participant can say what Beelo is—and is not—in their own words."),
        ("5–10 min", "Open the pilot URL, add to Home Screen and confirm supported browser/device.", "Beelo launches from the icon and returns after closing."),
        ("10–15 min", "Create and confirm the local passphrase; explain that forgotten passphrases cannot be recovered.", "Participant unlocks after a fresh launch without assistance."),
        ("15–20 min", "Complete profile, units, commission mode, map preference and permissions.", "Participant can find Settings and explain each permission chosen."),
        ("20–27 min", "Create a practice customer/visit, open navigation, mark Arrived, record an outcome and Leave.", "A complete visit journey appears in the correct records."),
        ("27–31 min", "Use Quick Add for a safe sample; review extracted values before saving.", "Participant corrects or rejects an inaccurate field."),
        ("31–34 min", "Show a communication draft and the WhatsApp/SMS hand-off.", "Participant states that they—not Beelo—review and send."),
        ("34–35 min", "Show backup, support and withdrawal route; agree first check-in.", "Participant knows where to get help and how to stop."),
    ], [1150, 4650, 3560], 8.25)
    add_heading(doc, "Onboarding acceptance", 2)
    add_checklist(doc, [
        "I launched and unlocked Beelo on my device.",
        "I understand that current working records are device-local and there is no cloud account recovery or automatic sync.",
        "I will verify scans, calculations, routes, dates and AI-generated text before acting on them.",
        "I understand that Beelo prepares drafts; I choose whether and how to send a message.",
        "I know how to create a safe backup and where I am allowed to store it.",
        "I know how to report a problem, withdraw, or ask for my pilot/research data to be handled according to the final notice.",
    ])
    add_field_line(doc, "Participant signature / date")
    add_field_line(doc, "Facilitator / date")


def demo(doc):
    page_break(doc)
    add_kicker(doc, "Founder / partner presentation")
    add_heading(doc, "Five-minute product demonstration", 1)
    add_callout(doc, "DEMO RULE", "Use a prepared local demo profile. Never expose a real participant's customer data, shared secret, API key, passphrase or personal notifications.")
    add_matrix(doc, ["Time", "Show", "Suggested words / proof point"], [
        ("0:00", "Opening problem", "“Seven tools, one person holding it all together. Beelo is the operational-memory layer for the person on the road.”"),
        ("0:35", "Home / My Day", "Show visits, due work and next context. Explain that Beelo surfaces information; it does not control the diary or reschedule silently."),
        ("1:15", "Quick Add", "Scan a safe sample and route it to a visit or expense. Correct one field deliberately to demonstrate human verification."),
        ("2:00", "Customer visit", "Open customer context, choose the map provider, mark Arrived, record time on site, outcome and next action, then Leave."),
        ("3:05", "Follow-up / message", "Open a due follow-up and prepare a draft. Pause on the review screen: “Beelo drafts. You decide.”"),
        ("3:55", "Order / job / money", "Show how commercial, operational and personal admin records remain linked without pretending Beelo is accounting or a supplier system."),
        ("4:35", "Pilot ask", "“We are not claiming measured impact yet. The pilot will test repeat use, mental load, safe AI control and fit for real field work.”"),
    ], [1000, 2200, 6160], 8.25)
    add_heading(doc, "Demo-day check", 2)
    add_checklist(doc, [
        "Demo profile is synthetic, complete and free of personal data.",
        "Device is charged; Do Not Disturb is on; screen recording and external display are tested.",
        "PWA, map hand-off, Quick Add, AI connection and offline fallback have been rehearsed.",
        "A screenshot/video fallback is available if a cloud dependency fails.",
        "Claims match the truth audit; voice, measured outcomes and team features are not implied.",
    ])


def daily_guide(doc):
    page_break(doc)
    add_kicker(doc, "Participant field card")
    add_heading(doc, "Using Beelo during a working day", 1)
    add_heading(doc, "Start of day", 2)
    add_label_para(doc, "REVIEW", "Open Home / My Day, check visits and due follow-ups, and confirm that the phone's date, time and notification settings are correct.")
    add_label_para(doc, "PREPARE", "Open the next customer record, review access/context and choose navigation. Treat route and arrival times as estimates.")
    add_heading(doc, "At the customer", 2)
    add_label_para(doc, "ARRIVE", "Use Arrived only when it is safe to do so. Beelo records on-site time; do not operate the phone while driving.")
    add_label_para(doc, "CAPTURE", "Add only authorised notes, photos or measurements. Avoid unnecessary personal or sensitive information.")
    add_label_para(doc, "LEAVE", "Record the outcome, next action and Leave. Check that the timer stopped and the visit is linked correctly.")
    add_heading(doc, "Between visits", 2)
    add_label_para(doc, "FOLLOW UP", "Review drafts before sending. Check recipient, dates, values and tone; opening WhatsApp/SMS does not prove delivery.")
    add_label_para(doc, "ADMIN", "Log a receipt, expense or mileage entry while the context is fresh. Verify extracted and calculated values.")
    add_heading(doc, "End of day / week", 2)
    add_label_para(doc, "CLOSE", "Complete missing outcomes and identify tomorrow's risks. Do not treat Beelo as the only source for critical commitments.")
    add_label_para(doc, "BACK UP", "Create the required encrypted backup on the agreed schedule, store it only in the approved place and test restore as instructed.")
    add_callout(doc, "15-MINUTE ALERTS", "Appointment alerts depend on device permission, operating-system behaviour and the app's running state. They are helpful prompts, not a guaranteed alarm. Keep the source diary and your own judgement in control.")


def safe_use(doc):
    page_break(doc)
    add_kicker(doc, "Participant safety")
    add_heading(doc, "Safe data and AI use", 1)
    add_matrix(doc, ["Do", "Avoid"], [
        ("Use the minimum customer information needed for the agreed workflow.", "Payment-card details, bank logins, identity documents, health data or other special-category data."),
        ("Check that you are authorised to store a customer photo, note, quote or supplier document.", "Copying company, franchise or supplier data when your agreement does not permit it."),
        ("Review every OCR field, calculation, route and AI draft before saving or acting.", "Assuming confident wording means the AI output is accurate."),
        ("Lock the phone, protect the Beelo passphrase and use an approved encrypted backup.", "Sharing passphrases, secrets or readable backup files by email or messaging."),
        ("Report unexpected disclosure, wrong-recipient risk, data loss or suspicious behaviour immediately.", "Trying repeated recovery actions that might overwrite the only copy of data."),
        ("Continue using the normal company/source system for authoritative bookings and orders.", "Treating Beelo as the sole legal, financial, tax, payment or supplier record."),
    ], [4680, 4680], 8.7)
    add_heading(doc, "How optional AI works", 2)
    add_para(doc, "When an AI feature is used, selected content needed for that request may be sent through Beelo's configured proxy to the AI provider. Core local workflows do not require AI, but AI-assisted OCR, drafting and phrasing require connectivity. The final participant notice must name the actual processors and transfer safeguards.")
    add_checklist(doc, [
        "I will use AI only for agreed low-risk tasks.",
        "I will remove unnecessary personal details before submitting content where practical.",
        "I will correct, reject or ignore any output that is inaccurate, inappropriate or uncertain.",
        "I understand that Beelo does not send a message for me and does not make a customer, schedule, financial or tax decision for me.",
    ])
    add_callout(doc, "STOP AND REPORT", "Pause use immediately if the app appears to expose another person's data, sends or changes something without approval, loses unrecoverable records, or behaves in a way that could materially harm a customer or participant.", PALE_RED, RED)


def privacy_consent(doc):
    page_break(doc)
    add_kicker(doc, "Review before issue")
    add_heading(doc, "Participant information and consent — operational draft", 1)
    add_callout(doc, "LEGAL / DATA-PROTECTION REVIEW REQUIRED", "Complete the bracketed items and obtain appropriate UK review before giving this page to participants. This draft does not determine the lawful basis or replace a full privacy notice.", PALE_RED, RED)
    add_heading(doc, "Why you are being invited", 2)
    add_para(doc, "You are invited to help evaluate an early Beelo prototype for solo field professionals. The purpose is to learn whether it is understandable, usable and helpful in real work, and to identify technical, privacy, safety and product improvements before any wider release.")
    add_heading(doc, "What participation involves", 2)
    add_para(doc, "Participation is expected to last about four weeks. It may include onboarding, agreed use during normal work, short weekly questionnaires, issue reports and a final interview. Participation is voluntary. You may pause or withdraw without having to give a reason, subject to the final notice explaining what can and cannot be removed from already anonymised analysis.")
    add_heading(doc, "Important prototype risks", 2)
    add_para(doc, "Beelo is an early prototype. It may contain defects, lose locally stored data, show inaccurate calculations or routes, miss notifications, or produce incorrect AI/OCR output. It is not an emergency, accounting, tax-filing, payment, legal or supplier-order system. Continue using your normal authoritative systems and professional judgement.")
    page_break(doc)
    add_heading(doc, "Data handling summary", 2)
    add_two_col_facts(doc, [
        ("Pilot operator", "[Legal entity, postal address, email and company number]"),
        ("Applicant data", "Name, contact details, trade, area and screening answers received through the pilot form/email."),
        ("Research data", "Onboarding record, questionnaires, interviews, support/issue reports and usage evidence agreed for the pilot."),
        ("Working data", "Customer, visit, job and money records entered into the participant's local Beelo installation."),
        ("Local storage", "Current Beelo working records are held in the browser/device; no account sync or automatic cloud recovery."),
        ("AI use", "Selected request content may pass through [proxy/hosting processor] to [AI provider] when the participant chooses an AI feature."),
        ("Retention", "[Set separate periods for unsuccessful applicants, participants, identifiable research, issue logs and backups]."),
        ("Rights / complaints", "[Explain rights, contact route and ICO complaint route based on reviewed lawful basis and arrangement]."),
    ])
    add_heading(doc, "Consent / acknowledgement", 2)
    add_checklist(doc, [
        "I have read the final participant information and had an opportunity to ask questions.",
        "I understand that participation is voluntary and how to pause or withdraw.",
        "I understand the prototype limitations, device-local storage and backup responsibilities.",
        "I agree to the research activities and data uses described in the final notice.",
        "I understand when selected content may be processed by the optional AI service.",
        "I agree to report problems and not enter prohibited or unauthorised information.",
    ])
    add_field_line(doc, "Participant name")
    add_field_line(doc, "Signature / date")
    add_field_line(doc, "Pilot lead / date")


def baseline(doc):
    page_break(doc)
    add_kicker(doc, "Research form")
    add_heading(doc, "Baseline questionnaire", 1)
    add_field_line(doc, "Participant ID / date")
    add_field_line(doc, "Role / trade / typical weekly field visits")
    add_field_line(doc, "Tools used today (diary, messaging, maps, photos, notes, receipts, mileage, spreadsheet, other)", 3)
    add_field_line(doc, "Three admin tasks that consume the most attention", 3)
    add_field_line(doc, "A recent example of forgotten context, duplication, delay or error", 3)
    add_heading(doc, "Starting ratings", 2)
    add_matrix(doc, ["Statement", "1", "2", "3", "4", "5"], [
        ("I can find the context I need before a visit", "□", "□", "□", "□", "□"),
        ("I reliably remember follow-ups", "□", "□", "□", "□", "□"),
        ("My mileage and expenses are captured consistently", "□", "□", "□", "□", "□"),
        ("I understand my order/job position", "□", "□", "□", "□", "□"),
        ("My current tool chain feels manageable", "□", "□", "□", "□", "□"),
    ], [5360, 800, 800, 800, 800, 800], 8.8)
    add_para(doc, "Scale: 1 = strongly disagree • 5 = strongly agree", size=8.5, color=GRAY, italic=True)
    add_field_line(doc, "Estimated admin hours in a typical week")
    add_field_line(doc, "What would make this pilot worthwhile for you?", 3)


def weekly(doc):
    page_break(doc)
    add_kicker(doc, "Research form")
    add_heading(doc, "Weekly pulse", 1)
    add_field_line(doc, "Participant ID / week / date")
    add_matrix(doc, ["This week…", "1", "2", "3", "4", "5", "N/A"], [
        ("Beelo made customer context easier to find", "□", "□", "□", "□", "□", "□"),
        ("Beelo reduced the need to remember tasks mentally", "□", "□", "□", "□", "□", "□"),
        ("I trusted myself to stay in control of drafts and actions", "□", "□", "□", "□", "□", "□"),
        ("The app was reliable enough for agreed pilot use", "□", "□", "□", "□", "□", "□"),
        ("I would choose to use Beelo next week", "□", "□", "□", "□", "□", "□"),
    ], [4560, 800, 800, 800, 800, 800, 800], 8.5)
    add_para(doc, "Scale: 1 = strongly disagree • 5 = strongly agree", size=8.5, color=GRAY, italic=True)
    add_field_line(doc, "Most useful moment this week", 3)
    add_field_line(doc, "Most frustrating or confusing moment", 3)
    add_field_line(doc, "Any wrong, risky or surprising result? What did you do?", 3)
    add_field_line(doc, "One change that would make next week better", 3)
    add_field_line(doc, "Approximate active days / visits / expenses / follow-ups used")


def interview(doc):
    page_break(doc)
    add_kicker(doc, "Research guide")
    add_heading(doc, "End-of-pilot interview", 1)
    add_para(doc, "Use these prompts conversationally. Ask for concrete examples before ratings or feature requests; avoid leading the participant toward the founder's preferred answer.")
    questions = [
        ("Story", "Talk me through the last working day when you used Beelo. Where did it fit naturally, and where did you work around it?"),
        ("Problem", "Which original admin problem mattered most? Did Beelo change it? What evidence makes you say that?"),
        ("Trust", "When did you trust a Beelo suggestion or draft? When did you correct, reject or avoid one?"),
        ("Control", "Did anything feel automatic or unclear? Were you always sure what would happen after tapping?"),
        ("Local-first", "How did you feel about device-local storage, the passphrase and backup responsibility?"),
        ("Habit", "What made you return—or forget to return—to Beelo? What would you keep using without a reminder?"),
        ("Replacement", "Which existing tool, if any, became less necessary? Which systems must remain authoritative?"),
        ("Value", "If Beelo disappeared tomorrow, what would you miss? Would you recommend it to someone like you? Why?"),
        ("Commercial", "What would have to be true before you would pay? What pricing approach would feel fair?"),
        ("Priority", "If only one journey could be improved next, which should it be and why?"),
    ]
    for label, q in questions[:7]:
        add_label_para(doc, label, q, after=3)
        add_field_line(doc, "Notes", 2)
    page_break(doc)
    add_kicker(doc, "Research guide")
    add_heading(doc, "End-of-pilot interview — continued", 1)
    for label, q in questions[7:]:
        add_label_para(doc, label, q, after=3)
        add_field_line(doc, "Notes", 2)
    add_heading(doc, "Final rating", 2)
    add_matrix(doc, ["Outcome", "Low", "Medium", "High", "Evidence / quote"], [
        ("Repeat-use intent", "□", "□", "□", ""),
        ("Problem importance", "□", "□", "□", ""),
        ("Usability", "□", "□", "□", ""),
        ("Trust / control", "□", "□", "□", ""),
        ("Commercial signal", "□", "□", "□", ""),
    ], [2200, 900, 1100, 900, 4260], 8.4)


def scorecard(doc):
    page_break(doc)
    add_kicker(doc, "Founder decision")
    add_heading(doc, "Pilot scorecard", 1)
    add_matrix(doc, ["Measure", "Target", "Actual", "Evidence / caveat"], [
        ("Eligible participants onboarded", "≥ 80% of accepted cohort", "", ""),
        ("Core setup + first visit completed", "≥ 80%", "", ""),
        ("Median active days after week one", "≥ 3 per week", "", ""),
        ("Participants reporting reduced mental load", "≥ 70%", "", ""),
        ("Participants intending continued use", "≥ 60%", "", ""),
        ("Safe backup completed", "100%", "", ""),
        ("High-severity unresolved incidents", "0", "", ""),
        ("Unintended / misunderstood message sending", "0", "", ""),
        ("Specific segment + repeatable need identified", "Yes / no", "", ""),
        ("Support effort sustainable for next cohort", "Yes / no", "", ""),
    ], [3000, 2050, 1350, 2960], 8.3)
    add_heading(doc, "Evidence quality", 2)
    add_matrix(doc, ["Evidence", "Count", "Confidence / limitation"], [
        ("Observed task completions", "", ""),
        ("Weekly pulses received", "", ""),
        ("Exit interviews completed", "", ""),
        ("Usage records reviewed with consent", "", ""),
        ("Partner / specialist review", "", ""),
    ], [3200, 1300, 4860], 8.8)
    add_field_line(doc, "Decision: GO / ITERATE / PAUSE")
    add_field_line(doc, "Reason and strongest evidence", 4)
    add_field_line(doc, "Next experiment, owner and date", 3)


def incidents(doc):
    page_break(doc)
    add_kicker(doc, "Support control")
    add_heading(doc, "Issues, incidents and stop criteria", 1)
    add_matrix(doc, ["Level", "Example", "Immediate response"], [
        ("P0 — stop", "Wrong-person disclosure, unintended communication/action, unrecoverable data loss affecting real work, suspected compromise", "Stop affected use; preserve evidence; protect people/data; notify pilot lead immediately; assess external obligations."),
        ("P1 — urgent", "Core journey blocked with material work impact or repeated corruption risk", "Pause affected journey; safe workaround only; same-day owner and decision."),
        ("P2 — normal", "Important defect with workaround; confusing state; missed non-critical reminder", "Log, reproduce and prioritise for controlled release."),
        ("P3 — minor", "Cosmetic issue or low-impact improvement", "Record for review; do not interrupt pilot unnecessarily."),
    ], [1200, 4060, 4100], 8.1)
    add_heading(doc, "Issue record", 2)
    add_field_line(doc, "Issue ID / participant ID / date and time")
    add_field_line(doc, "Device / OS / browser / app release")
    add_field_line(doc, "Screen and action immediately before issue", 2)
    add_field_line(doc, "What happened / expected result", 3)
    add_field_line(doc, "Data involved and potential impact", 3)
    add_field_line(doc, "Reproduction steps / screenshots / logs", 3)
    add_field_line(doc, "Severity / workaround / owner / status", 2)
    page_break(doc)
    add_kicker(doc, "Support control")
    add_heading(doc, "Controlled-fix and release check", 1)
    add_checklist(doc, [
        "Participant data is protected or backed up before troubleshooting.",
        "The defect is reproduced against the named pilot release or documented as environment-specific.",
        "Source and minified production bundles both contain the fix; tests and build checks pass.",
        "The change is verified on the affected viewport/device and does not introduce a truth-claim change.",
        "Participants receive a clear update only if action is required; release and evidence records are updated.",
    ])
    add_heading(doc, "Release record", 2)
    add_field_line(doc, "Issue / fix ID and owner")
    add_field_line(doc, "Source commit / build / deployment")
    add_field_line(doc, "Tests and device evidence", 2)
    add_field_line(doc, "Participant communication / action required", 2)
    add_field_line(doc, "Verified by / date")


def closeout(doc):
    page_break(doc)
    add_kicker(doc, "Pilot completion")
    add_heading(doc, "Closeout checklist and findings memo", 1)
    add_checklist(doc, [
        "Complete the final interview and record permission for any attributed or anonymised quote.",
        "Resolve the participant's choice to continue, pause, export or remove pilot working data according to the reviewed notice.",
        "Create and verify the final approved backup; do not ask the participant to email an unencrypted operational backup.",
        "Close or transfer outstanding support issues and communicate known limitations.",
        "Separate measured results from founder observations, application interest and partner opinion.",
        "Compare baseline, weekly and final evidence; record missing responses and selection bias.",
        "Make the go / iterate / pause decision and identify the next smallest experiment.",
        "Thank the participant and state what feedback changed; do not imply a commercial launch date unless decided.",
    ])
    add_heading(doc, "One-page findings memo", 2)
    add_field_line(doc, "Pilot question and cohort")
    add_field_line(doc, "What we observed (facts only)", 3)
    add_field_line(doc, "What participants said (with consent / attribution status)", 3)
    add_field_line(doc, "What remains unknown", 3)
    add_field_line(doc, "Safety, privacy and reliability findings", 3)
    add_field_line(doc, "Decision and next experiment", 3)


def claims_card(doc):
    page_break(doc)
    add_kicker(doc, "Meeting leave-behind")
    add_heading(doc, "Beelo truth and claims card", 1)
    add_matrix(doc, ["Safe to say", "Add the boundary"], [
        ("“Beelo is an installable, offline-capable PWA for a solo field advisor.”", "Core local workflows are designed for offline use; cloud maps, weather and AI can degrade or require connectivity."),
        ("“Working records are stored locally on the device.”", "There is no cloud sync, server backup or account recovery in the current prototype."),
        ("“Beelo prepares messages for review.”", "The advisor edits and decides whether to hand off to WhatsApp/SMS; delivery is not inferred."),
        ("“AI is optional and human-controlled.”", "It requires connectivity, may be wrong and transmits selected context through the configured proxy/provider."),
        ("“The product was built and founder-tested from lived field experience.”", "Independent pilot evidence is being collected; do not call this validated demand."),
        ("“Visits, follow-ups, orders, jobs, expenses and mileage are implemented.”", "They do not replace company, supplier, accounting or tax systems."),
    ], [4500, 4860], 8.5)
    add_heading(doc, "Do not claim yet", 2)
    add_checklist(doc, [
        "Proven reduction in admin time, errors, travel or missed follow-ups.",
        "Validated demand, commercial scalability or willingness to pay.",
        "Production-ready, fully offline, secure, compliant or independently certified.",
        "Voice capture, team accounts, manager dashboards, cloud sync or automatic recovery.",
        "Guaranteed notifications, routes, OCR, AI output, tax outcomes or message delivery.",
    ])
    add_callout(doc, "FOUNDER LINE", "“I didn't need another app. I needed the apps I already used to stop making me remember everything.” — Riaz, window-coverings advisor and founder of Beelo")


def invitation(doc):
    page_break(doc)
    add_kicker(doc, "Participant communication")
    add_heading(doc, "Pilot invitation template", 1)
    add_para(doc, "Subject: Help test Beelo — a four-week pilot for solo field professionals", size=10.5, bold=True, after=14)
    add_para(doc, "Hello [Name],")
    add_para(doc, "I am inviting a small group of UK-based field professionals to help test Beelo, an early device-local app built for the person managing visits, customer context, follow-ups and admin while working alone.")
    add_para(doc, "The pilot is not a public launch. It is a controlled four-week learning programme designed to test whether Beelo is understandable, reliable enough for agreed use, and genuinely helpful during real work. You would receive a short onboarding session, use the app for agreed low-risk workflows, complete brief weekly feedback and join a final interview.")
    add_para(doc, "Beelo does not replace your company diary, supplier, accounting or tax systems. Its current records are stored on your device, and optional AI features require your review and approval. We will explain the limitations, privacy arrangements and withdrawal route before you decide whether to take part.")
    add_para(doc, "If you are interested, please reply with a suitable time for a short eligibility conversation. There is no obligation to join.")
    add_para(doc, "Kind regards,\nMuhammad Asif Riaz\nFounder, Beelo\nhello@beelestial.co.uk")
    add_heading(doc, "Issue history", 2)
    add_two_col_facts(doc, [
        ("v1.0 — 25 Aug 2026", "Initial controlled-pilot pack created from the Beelo truth audit and current pilot preparation work."),
        ("Next review", "After operator/privacy details are completed and before the first external participant receives the pack."),
    ])
    add_para(doc, "End of pack", size=9, color=GRAY, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=30)


def build():
    doc = Document()
    setup_styles(doc)
    cover(doc)
    contents(doc)
    pilot_at_glance(doc)
    hypotheses_and_success(doc)
    readiness(doc)
    roles(doc)
    recruitment(doc)
    onboarding(doc)
    demo(doc)
    daily_guide(doc)
    safe_use(doc)
    privacy_consent(doc)
    baseline(doc)
    weekly(doc)
    interview(doc)
    scorecard(doc)
    incidents(doc)
    closeout(doc)
    claims_card(doc)
    invitation(doc)
    setup_header_footer(doc)

    props = doc.core_properties
    props.title = "Beelo Pilot Pack"
    props.subject = "Controlled UK pilot operations, onboarding and evaluation pack"
    props.creator = "Beelo"
    props.last_modified_by = "Beelo"
    props.keywords = "Beelo, pilot, field professionals, onboarding, evaluation"
    props.comments = "Operational working document; privacy and consent sections require review before external use."

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
